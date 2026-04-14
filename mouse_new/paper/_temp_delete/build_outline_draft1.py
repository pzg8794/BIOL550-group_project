#!/usr/bin/env python3
from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    style_run(run)


PAPER_DIR = Path(__file__).resolve().parent
MD_PATH = PAPER_DIR / "HTSA_Paper_Outline.md"
TEX_PATH = PAPER_DIR / "HTSA_Paper_Outline.tex"
DOCX_PATH = PAPER_DIR / "HTSA_Paper_Outline.docx"
PDF_PATH = PAPER_DIR / "HTSA_Paper_Outline.pdf"


@dataclass(frozen=True)
class Contributor:
    key: str
    name: str
    color_name: str
    color_hex: str
    fill_hex: str
    initials: str
    ownership: str


@dataclass(frozen=True)
class Bullet:
    text: str
    owner: str | None = None
    children: tuple["Bullet", ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class Section:
    title: str
    owner: str | None
    bullets: tuple[Bullet, ...]


CONTRIBUTORS = {
    "nikhi": Contributor(
        key="nikhi",
        name="Nikhi Boggavarapu",
        color_name="light purple",
        color_hex="8E7CC3",
        fill_hex="D9D2E9",
        initials="NB",
        ownership="Introduction, Discussion, and back matter",
    ),
    "sam": Contributor(
        key="sam",
        name="Sam Kopelev",
        color_name="light green",
        color_hex="6AA84F",
        fill_hex="B6D7A8",
        initials="SK",
        ownership="Results framing and computational/results interpretation",
    ),
    "piter": Contributor(
        key="piter",
        name="Piter Garcia",
        color_name="light red",
        color_hex="CC6666",
        fill_hex="F4CCCC",
        initials="PG",
        ownership="Materials and Methods and statistical/analysis framing",
    ),
}


def bullet(text: str, *children: Bullet, owner: str | None = None) -> Bullet:
    return Bullet(text=text, owner=owner, children=tuple(children))


SECTIONS = (
    Section(
        title="Introduction",
        owner="nikhi",
        bullets=(
            bullet(
                "Brief intro about high-throughput sequencing",
                bullet("RNA-seq as a transcriptome-wide method for detecting gene expression changes", owner="piter"),
                bullet(
                    "Types of NGS",
                    bullet("Organized by read length"),
                    bullet("Applications"),
                    owner="nikhi",
                ),
                bullet(
                    "Benefits of NGS",
                    bullet("High throughput"),
                    bullet("Fast"),
                    owner="nikhi",
                ),
                bullet(
                    "What our goal was in utilizing NGS",
                    bullet("Unbiased, transcriptome-wide detection of gene expression changes", owner="piter"),
                    owner="nikhi",
                ),
                owner="nikhi",
            ),
            bullet(
                "Introduce paper/dataset",
                bullet("Experimental design"),
                bullet("DRG after sciatic nerve injury"),
                bullet(
                    "Goals of the paper",
                    bullet("Why did they analyze DRG with cKO"),
                    bullet("Ipsilateral vs contralateral"),
                    bullet("FF / cre context", owner="piter"),
                    owner="sam",
                ),
                owner="sam",
            ),
            bullet(
                "How we wanted to use the paper",
                bullet(
                    "Explore and confirm suggested pathways",
                    bullet("FF / cre"),
                    owner="sam",
                ),
                bullet("Examine gene expression changes beyond the genes emphasized in the original paper"),
                owner="sam",
            ),
            bullet(
                "Claim",
                bullet("*Differential expression (DE):* transcriptome-wide RNA-seq helps identify genes involved beyond the ones the paper discussed"),
                bullet("*Key contrasts:* expression differences were quantified across the main experimental comparisons, especially ipsilateral vs contralateral DRG, and genotype where relevant"),
                bullet("*GO / pathway enrichment:* DE gene sets were interpreted in terms of broader biological processes and pathways linked to injury response and the cKO context"),
                owner="piter",
            ),
        ),
    ),
    Section(
        title="Materials and Methods",
        owner="piter",
        bullets=(
            bullet(
                "Dataset and study design",
                bullet("Public dataset accession"),
                bullet("Tissue and injury context"),
                bullet("Sample groups and contrasts"),
                bullet("Why was this subset retained"),
                owner="piter",
            ),
            bullet(
                "Pipeline Steps",
                bullet(
                    "Data Collection and Preprocessing",
                    bullet("Download SRR files (`.gz`)", owner="nikhi"),
                    bullet("FASTQ organization", owner="piter"),
                    bullet("Sample manifest / metadata table", owner="piter"),
                    bullet("Tools used", owner="piter"),
                    owner="nikhi",
                ),
                bullet(
                    "Data Quality Control and Trimming",
                    bullet("Raw FastQC / MultiQC"),
                    bullet("FASTP trimming"),
                    bullet("Post-trim FastQC / MultiQC"),
                    bullet("Key QC checkpoint metrics"),
                    owner="nikhi",
                ),
                bullet(
                    "Data Preparation (Alignment and count generation)",
                    bullet(
                        "Reference genome and annotation",
                        bullet("(mm39)"),
                        owner="nikhi",
                    ),
                    bullet(
                        "STAR index and alignment",
                        bullet("Building the STAR genome"),
                        bullet("Paired-end alignment"),
                        owner="nikhi",
                    ),
                    bullet("GeneCounts output", owner="piter"),
                    bullet("BAM / log / count outputs", owner="piter"),
                    bullet("Family count matrix assembly", owner="piter"),
                    owner="nikhi",
                ),
                bullet(
                    "Data Analysis and Interpretation",
                    bullet("MultiQC", owner="sam"),
                    bullet(
                        "Differential Expression Analysis",
                        bullet("DESeq2 setup"),
                        bullet(
                            "DESeq2",
                            bullet("Volcano plot"),
                            bullet("Heatmap of top DEGs (Regeneration-Enhancing)"),
                            bullet("MA plot", owner="piter"),
                            bullet("Target validation (gene expression boxplots)"),
                            owner="sam",
                        ),
                        bullet("Filtering rule", owner="piter"),
                        bullet("Modeled contrasts", owner="piter"),
                        bullet("Thresholds / significance criteria", owner="piter"),
                        owner="sam",
                    ),
                    bullet(
                        "Follow-up selection for functional interpretation",
                        bullet("Bend-point rule", owner="piter"),
                        bullet("GO analysis", owner="sam"),
                        bullet("GSEA analysis", owner="sam"),
                        bullet("Panther analysis", owner="sam"),
                        bullet("Extra analysis, if finalized", owner="sam"),
                        owner="sam",
                    ),
                    owner="sam",
                ),
                owner="nikhi",
            ),
        ),
    ),
    Section(
        title="Results",
        owner="sam",
        bullets=(
            bullet(
                "Dataset quality supported downstream analysis",
                bullet("Number of samples and read structure"),
                bullet(
                    "Pre-trim",
                    bullet("GC content"),
                    bullet("Adapter sequences"),
                    owner="sam",
                ),
                bullet(
                    "Alignment stats",
                    bullet("Unique mapping"),
                    owner="sam",
                ),
                owner="sam",
            ),
            bullet(
                "The sample structure showed the strongest separation by the main biological contrast",
                bullet("PCA (key figure)"),
                bullet("Interpretation of PC1 / PC2"),
                bullet("Genotype analysis"),
                owner="piter",
            ),
            bullet(
                "Differential expression identified the strongest transcriptomic changes",
                bullet(
                    "PCA (key figure)",
                    bullet("Explain how it is separated, PC1 and PC2"),
                    owner="sam",
                ),
                bullet("MA plot", owner="nikhi"),
                bullet("Cumulative Distribution Plot (bend-point / elbow rule)", owner="nikhi"),
                bullet(
                    "Volcano Plot (Important Genes)",
                    bullet("Include a secondary Volcano Plot with a new threshold", owner="nikhi"),
                    owner="nikhi",
                ),
                bullet(
                    "Heatmap",
                    bullet("Including the interpretation of the distance heatmap between the two conditions"),
                    owner="nikhi",
                ),
                owner="sam",
            ),
            bullet(
                "Bend-point selection narrowed the main follow-up sets",
                bullet("Cumulative plot"),
                bullet("Selected gene counts"),
                bullet("Why this helped interpretation"),
                bullet("Secondary volcano plot, if you keep it"),
                owner="piter",
            ),
            bullet(
                "Functional enrichment connected DE results to broader biology",
                bullet("GO / pathway results"),
                bullet(
                    "Pathway level / Panther analysis",
                    bullet("Go through the plots listed by Panther"),
                    owner="nikhi",
                ),
                bullet("Proteostasis"),
                bullet("Translation"),
                bullet("Metabolism"),
                bullet("Extra analysis, once finalized"),
                owner="nikhi",
            ),
        ),
    ),
    Section(
        title="Discussion",
        owner="nikhi",
        bullets=(
            bullet(
                "Main biological interpretation (what does the data show)",
                bullet("What the data show overall"),
                bullet("Strongest supported signal"),
                bullet("What is primary vs secondary"),
                bullet(
                    "Gene expression differences",
                    bullet("Injury vs control"),
                    owner="piter",
                ),
                bullet(
                    "Pathways being affected",
                    bullet("Proteostasis (AhR activation)"),
                    bullet("Translation (suppression / upregulation of genes)"),
                    bullet("Metabolism (energy output)"),
                    owner="piter",
                ),
                owner="piter",
            ),
            bullet(
                "What was unusual or needs caution (things that were weird about the dataset to consider)",
                bullet("Odd volcano plots"),
                bullet("The genotype signal appears weaker / secondary"),
                bullet("PCA collisions / overlap"),
                bullet("Interpretation vs causation"),
                owner="nikhi",
            ),
            bullet(
                "What this adds beyond the original paper",
                bullet(
                    "Expands the analysis beyond the paper’s candidate genes",
                    bullet(
                        "Helps link pathways",
                        bullet("Proposed and new"),
                        owner="sam",
                    ),
                    owner="piter",
                ),
                bullet("Helps link the pathways already proposed in the paper with additional transcriptomic signals"),
                bullet("Connects newly identified genes to the genes highlighted in the paper through broader transcriptomic patterns"),
                bullet("Moves interpretation from single-gene emphasis toward pathway- and network-level context"),
                owner="piter",
            ),
            bullet(
                "NGS in the context of this paper",
                bullet("Global discovery"),
                bullet("Reproducibility"),
                bullet(
                    "External applications / biological relevance",
                    bullet("How does our global discovery connect to the genes found in the paper"),
                    bullet("Two-hybrid screening"),
                    owner="sam",
                ),
                owner="sam",
            ),
            bullet(
                "Future validation (what it suggests for future validation)",
                bullet("qPCR for key differentially expressed genes"),
                bullet("Functional follow-up experiments for injury-response candidates"),
                bullet("Two-hybrid screening only if it remains biologically justified"),
                owner="piter",
            ),
            bullet(
                "Limitations and cautions",
                bullet("Weaker genotype signal"),
                bullet("PCA overlap / collisions"),
                bullet("Dependence on the original dataset design"),
                bullet("Interpretation vs causation"),
                owner="piter",
            ),
        ),
    ),
    Section(
        title="References (APA)",
        owner=None,
        bullets=(
            bullet("[Main Paper](https://www.nature.com/articles/s41586-026-10295-z#code-availability)"),
            bullet("NGS overview / background sources"),
            bullet("Additional papers using or reviewing NGS"),
            bullet("Background sources for the mouse study"),
            bullet("Minimum of 15 primary sources in the final paper"),
        ),
    ),
    Section(
        title="Code Availability",
        owner=None,
        bullets=(
            bullet("Git repository"),
            bullet("Notebook(s)"),
        ),
    ),
    Section(
        title="Source Data",
        owner=None,
        bullets=(
            bullet("CSVs / tables supporting plots"),
            bullet("Data frames / exported tables"),
        ),
    ),
    Section(
        title="Supplemental Material",
        owner=None,
        bullets=(
            bullet("Full figure set"),
            bullet("Additional supporting outputs"),
        ),
    ),
)


TITLE = "HTSA Group Paper Outline"
SUBTITLE = "BIOL550 Group Project (Draft 1)"
AUTHORS = ("Nikhi Boggavarapu", "Sam Kopelev", "Piter Garcia")


def escape_tex(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = []
    for ch in text:
        out.append(replacements.get(ch, ch))
    return "".join(out)


def render_inline_markdown_to_tex(text: str) -> str:
    if text.startswith("[") and "](" in text and text.endswith(")"):
        label, url = text[1:].split("](", 1)
        return rf"\href{{{escape_tex(url[:-1])}}}{{{escape_tex(label)}}}"
    if text.startswith("*") and ":* " in text:
        label, remainder = text.split(":* ", 1)
        return rf"\textit{{{escape_tex(label[1:] + ':')}}} {escape_tex(remainder)}"
    if text.startswith("*") and text.endswith("*"):
        return rf"\textit{{{escape_tex(text[1:-1])}}}"
    return escape_tex(text)


def render_inline_markdown_to_docx(paragraph, text: str, color_hex: str | None = None) -> list:
    runs = []
    rgb = RGBColor.from_string(color_hex) if color_hex else None
    if text.startswith("[") and "](" in text and text.endswith(")"):
        label = text[1:].split("](", 1)[0]
        run = paragraph.add_run(label)
        run.underline = True
        if rgb:
            run.font.color.rgb = rgb
        runs.append(run)
        return runs
    if text.startswith("*") and ":* " in text:
        label, remainder = text.split(":* ", 1)
        run1 = paragraph.add_run(label[1:] + ": ")
        run1.italic = True
        if rgb:
            run1.font.color.rgb = rgb
        run2 = paragraph.add_run(remainder)
        if rgb:
            run2.font.color.rgb = rgb
        runs.extend([run1, run2])
        return runs
    run = paragraph.add_run(text)
    if rgb:
        run.font.color.rgb = rgb
    runs.append(run)
    return runs


def markdown_lines() -> list[str]:
    lines = [*AUTHORS, "", f"**{TITLE}**", f"**{SUBTITLE}**", ""]
    for key in ("nikhi", "sam", "piter"):
        contributor = CONTRIBUTORS[key]
        lines.append(f"* {contributor.name}")
    lines.append("")
    for section in SECTIONS:
        lines.append(f"**{section.title}**")
        lines.append("")
        lines.extend(render_markdown_bullets(section.bullets, 0))
        lines.append("")
    return lines


def render_markdown_bullets(bullets: Iterable[Bullet], level: int) -> list[str]:
    lines: list[str] = []
    indent = "  " * level
    for item in bullets:
        lines.append(f"{indent}* {item.text}")
        if item.children:
            lines.extend(render_markdown_bullets(item.children, level + 1))
    return lines


def build_markdown() -> None:
    MD_PATH.write_text("\n".join(markdown_lines()).rstrip() + "\n")


def render_tex_bullets(bullets: Iterable[Bullet], owner: Contributor | None, level: int = 0) -> list[str]:
    lines = [r"\begin{itemize}"]
    for item in bullets:
        item_owner = CONTRIBUTORS[item.owner] if item.owner else owner
        if item_owner is None:
            lines.append(rf"\item {render_inline_markdown_to_tex(item.text)}")
        else:
            lines.append(rf"\item \textcolor{{{item_owner.key}color}}{{{render_inline_markdown_to_tex(item.text)}}}")
        if item.children:
            lines.extend(render_tex_bullets(item.children, item_owner, level + 1))
    lines.append(r"\end{itemize}")
    return lines


def build_tex() -> None:
    lines = [
        r"\documentclass[12pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{setspace}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{newtxtext,newtxmath}",
        r"\usepackage{xcolor}",
        r"\usepackage{colortbl}",
        r"\usepackage{enumitem}",
        r"\usepackage{hyperref}",
        r"\doublespacing",
        r"\setlength{\parindent}{0pt}",
        r"\setlength{\parskip}{0.2em}",
        r"\setlistdepth{6}",
        r"\renewlist{itemize}{itemize}{6}",
        r"\setlist[itemize,1]{label=\textbullet,leftmargin=1.6em,itemsep=0.15em,topsep=0.15em}",
        r"\setlist[itemize,2]{label=\textbullet,leftmargin=1.9em,itemsep=0.15em,topsep=0.15em}",
        r"\setlist[itemize,3]{label=\textbullet,leftmargin=2.2em,itemsep=0.15em,topsep=0.15em}",
        r"\setlist[itemize,4]{label=\textbullet,leftmargin=2.5em,itemsep=0.15em,topsep=0.15em}",
        r"\setlist[itemize,5]{label=\textbullet,leftmargin=2.8em,itemsep=0.15em,topsep=0.15em}",
        r"\setlist[itemize,6]{label=\textbullet,leftmargin=3.1em,itemsep=0.15em,topsep=0.15em}",
        rf"\definecolor{{nikhicolor}}{{HTML}}{{{CONTRIBUTORS['nikhi'].color_hex}}}",
        rf"\definecolor{{samcolor}}{{HTML}}{{{CONTRIBUTORS['sam'].color_hex}}}",
        rf"\definecolor{{pitercolor}}{{HTML}}{{{CONTRIBUTORS['piter'].color_hex}}}",
        r"\begin{document}",
        r"\begin{center}",
        rf"{{\Large \textbf{{{escape_tex(TITLE)}}}}}\\",
        rf"{{\normalsize {escape_tex(SUBTITLE)}}}\\[0.5em]",
    ]
    for name in AUTHORS:
        lines.append(rf"{{\normalsize {escape_tex(name)}}}\\")
    lines.extend(
        [
            r"\end{center}",
            r"\vspace{0.5em}",
            r"\noindent",
            r"\begin{tabular}{|p{1.55in}|p{4.65in}|}",
            r"\hline",
            r"\textbf{Member} & \textbf{Note} \\ \hline",
        ]
    )
    for key in ("nikhi", "sam", "piter"):
        contributor = CONTRIBUTORS[key]
        lines.append(
            rf"\cellcolor[HTML]{{{contributor.fill_hex}}}{escape_tex(contributor.name)} & "
            rf"\cellcolor[HTML]{{{contributor.fill_hex}}}Highlighted text in {escape_tex(contributor.color_name)} indicates {escape_tex(contributor.name.split()[0])}'s contribution. \\ \hline"
        )
    lines.append(r"\end{tabular}")
    for section in SECTIONS:
        owner = CONTRIBUTORS[section.owner] if section.owner else None
        if owner is None:
            lines.append(rf"\section*{{{escape_tex(section.title)}}}")
        else:
            lines.append(rf"\section*{{\textcolor{{{owner.key}color}}{{{escape_tex(section.title)}}}}}")
        lines.extend(render_tex_bullets(section.bullets, owner))
    lines.append(r"\end{document}")
    TEX_PATH.write_text("\n".join(lines) + "\n")


def apply_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches: float) -> None:
    width = Inches(inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def force_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def style_run(run, *, size: int = 12, bold: bool = False, italic: bool = False, color_hex: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def set_run_shading(run, fill_hex: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def add_heading(doc: Document, text: str, owner: Contributor | None) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.clear()
    run = p.add_run(text)
    style_run(run, size=13, bold=True)


def add_bullet_paragraph(doc: Document, text: str, owner: Contributor, level: int) -> None:
    style_name = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    extra_indent = max(level - 2, 0)
    if extra_indent:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    runs = render_inline_markdown_to_docx(p, text, None)
    for run in runs:
        style_run(run)
        if owner.fill_hex:
            set_run_shading(run, owner.fill_hex)


def add_bullets(doc: Document, bullets: Iterable[Bullet], owner: Contributor | None, level: int = 0) -> None:
    for item in bullets:
        item_owner = CONTRIBUTORS[item.owner] if item.owner else owner
        if item_owner is None:
            plain_owner = Contributor("plain", "", "", "000000", "", "", "")
            add_bullet_paragraph(doc, item.text, plain_owner, level)
        else:
            add_bullet_paragraph(doc, item.text, item_owner, level)
        if item.children:
            add_bullets(doc, item.children, item_owner, level + 1)


def build_docx() -> None:
    doc = Document()
    apply_doc_defaults(doc)
    footer_para = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer_para)

    for text, size, bold in (
        (TITLE, 16, True),
        (SUBTITLE, 12, False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        style_run(run, size=size, bold=bold)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    authors.paragraph_format.space_after = Pt(10)
    for idx, name in enumerate(AUTHORS):
        if idx:
            authors.add_run(" | ")
        run = authors.add_run(name)
        style_run(run)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    force_fixed_table_layout(table)
    left_col_width = 1.55
    right_col_width = 4.65
    table.columns[0].width = Inches(left_col_width)
    table.columns[1].width = Inches(right_col_width)
    for row in table.rows:
        set_cell_width(row.cells[0], left_col_width)
        set_cell_width(row.cells[1], right_col_width)
    hdr0 = table.cell(0, 0)
    hdr1 = table.cell(0, 1)
    hdr0.text = ""
    hdr1.text = ""
    p0 = hdr0.paragraphs[0]
    p1 = hdr1.paragraphs[0]
    r0 = p0.add_run("Member")
    r1 = p1.add_run("Note")
    style_run(r0, bold=True)
    style_run(r1, bold=True)
    for row_idx, key in enumerate(("nikhi", "sam", "piter"), start=1):
        contributor = CONTRIBUTORS[key]
        c0 = table.cell(row_idx, 0)
        c1 = table.cell(row_idx, 1)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run0 = p0.add_run(contributor.name)
        run1 = p1.add_run(
            f"Highlighted text in {contributor.color_name} indicates {contributor.name.split()[0]}'s contribution."
        )
        style_run(run0)
        style_run(run1)
        set_cell_shading(c0, contributor.fill_hex)
        set_cell_shading(c1, contributor.fill_hex)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    for section in SECTIONS:
        owner = CONTRIBUTORS[section.owner] if section.owner else None
        add_heading(doc, section.title, owner)
        add_bullets(doc, section.bullets, owner)

    doc.save(DOCX_PATH)


def render_pdf() -> None:
    if shutil.which("soffice") is None:
        raise RuntimeError("LibreOffice (soffice) is required to render the outline PDF.")
    profile_dir = Path(tempfile.mkdtemp(prefix="lo_profile_outline_"))
    try:
        subprocess.run(
            [
                shutil.which("soffice"),
                f"-env:UserInstallation=file://{profile_dir}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(PAPER_DIR),
                str(DOCX_PATH),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


def main() -> None:
    build_markdown()
    build_tex()
    build_docx()
    render_pdf()
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {TEX_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()
